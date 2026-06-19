from __future__ import annotations

import csv
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION_START
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


BASE_DIR = Path(__file__).resolve().parents[1]
OUTPUT = BASE_DIR / "Informe_TPO_Comportamiento_Parciales.docx"

BLUE = RGBColor(46, 116, 181)
DARK_BLUE = RGBColor(31, 77, 120)
INK = RGBColor(30, 41, 59)
MUTED = RGBColor(100, 116, 139)
LIGHT_FILL = "F2F4F7"
CALLOUT_FILL = "E8EEF5"


def read_csv(path: Path) -> list[dict[str, str]]:
    if not path.exists():
        return []
    with path.open("r", encoding="utf-8-sig", newline="") as handle:
        return list(csv.DictReader(handle))


def output_csv(name: str) -> Path:
    output_path = BASE_DIR / "outputs" / name
    return output_path if output_path.exists() else BASE_DIR / name


def fmt_num(value: str | float | None, decimals: int = 2) -> str:
    if value in (None, ""):
        return "-"
    try:
        return f"{float(value):.{decimals}f}"
    except (TypeError, ValueError):
        return str(value)


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=80, start=120, bottom=80, end=120) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for margin, value in {
        "top": top,
        "start": start,
        "bottom": bottom,
        "end": end,
    }.items():
        node = tc_mar.find(qn(f"w:{margin}"))
        if node is None:
            node = OxmlElement(f"w:{margin}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_table_width(table, widths: list[int]) -> None:
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    tbl = table._tbl
    tbl_pr = tbl.tblPr
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(sum(widths)))
    tbl_w.set(qn("w:type"), "dxa")

    tbl_grid = tbl.tblGrid
    for child in list(tbl_grid):
        tbl_grid.remove(child)
    for width in widths:
        grid_col = OxmlElement("w:gridCol")
        grid_col.set(qn("w:w"), str(width))
        tbl_grid.append(grid_col)

    for row in table.rows:
        for idx, cell in enumerate(row.cells):
            cell.width = Pt(widths[idx] / 20)
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_w = tc_pr.find(qn("w:tcW"))
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                tc_pr.append(tc_w)
            tc_w.set(qn("w:w"), str(widths[idx]))
            tc_w.set(qn("w:type"), "dxa")
            set_cell_margins(cell)
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER


def style_table(table, widths: list[int]) -> None:
    table.style = "Table Grid"
    set_table_width(table, widths)
    for row_idx, row in enumerate(table.rows):
        for cell in row.cells:
            for paragraph in cell.paragraphs:
                paragraph.style = "Table Text"
            if row_idx == 0:
                set_cell_shading(cell, LIGHT_FILL)
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        run.bold = True
                        run.font.color.rgb = INK


def add_table(doc: Document, headers: list[str], rows: list[list[str]], widths: list[int]):
    table = doc.add_table(rows=1, cols=len(headers))
    header_cells = table.rows[0].cells
    for idx, header in enumerate(headers):
        header_cells[idx].text = header
    for row_values in rows:
        cells = table.add_row().cells
        for idx, value in enumerate(row_values):
            cells[idx].text = value
    style_table(table, widths)
    return table


def add_callout(doc: Document, title: str, body: str) -> None:
    table = doc.add_table(rows=1, cols=1)
    cell = table.rows[0].cells[0]
    set_cell_shading(cell, CALLOUT_FILL)
    set_cell_margins(cell, top=140, bottom=140, start=180, end=180)
    p = cell.paragraphs[0]
    p.style = "Callout"
    run = p.add_run(title)
    run.bold = True
    run.font.color.rgb = DARK_BLUE
    p.add_run(f" {body}")
    set_table_width(table, [9360])
    doc.add_paragraph()


def add_bullet(doc: Document, text: str) -> None:
    p = doc.add_paragraph(style="List Bullet")
    p.add_run(text)


def add_number(doc: Document, text: str) -> None:
    p = doc.add_paragraph(style="List Number")
    p.add_run(text)


def configure_document(doc: Document) -> None:
    section = doc.sections[0]
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(11)
    normal.font.color.rgb = INK
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.10

    for style_name in ["Heading 1", "Heading 2", "Heading 3"]:
        style = styles[style_name]
        style.font.name = "Calibri"
        style.font.bold = True
    styles["Heading 1"].font.size = Pt(16)
    styles["Heading 1"].font.color.rgb = BLUE
    styles["Heading 1"].paragraph_format.space_before = Pt(16)
    styles["Heading 1"].paragraph_format.space_after = Pt(8)
    styles["Heading 2"].font.size = Pt(13)
    styles["Heading 2"].font.color.rgb = BLUE
    styles["Heading 2"].paragraph_format.space_before = Pt(12)
    styles["Heading 2"].paragraph_format.space_after = Pt(6)
    styles["Heading 3"].font.size = Pt(12)
    styles["Heading 3"].font.color.rgb = DARK_BLUE
    styles["Heading 3"].paragraph_format.space_before = Pt(8)
    styles["Heading 3"].paragraph_format.space_after = Pt(4)

    for style_name in ["List Bullet", "List Number"]:
        style = styles[style_name]
        style.font.name = "Calibri"
        style.font.size = Pt(11)
        style.paragraph_format.left_indent = Inches(0.5)
        style.paragraph_format.first_line_indent = Inches(-0.25)
        style.paragraph_format.space_after = Pt(6)
        style.paragraph_format.line_spacing = 1.167

    if "Table Text" not in styles:
        styles.add_style("Table Text", 1)
    table_text = styles["Table Text"]
    table_text.font.name = "Calibri"
    table_text.font.size = Pt(9.5)
    table_text.font.color.rgb = INK
    table_text.paragraph_format.space_after = Pt(0)
    table_text.paragraph_format.line_spacing = 1.10

    if "Callout" not in styles:
        styles.add_style("Callout", 1)
    callout = styles["Callout"]
    callout.font.name = "Calibri"
    callout.font.size = Pt(10.5)
    callout.font.color.rgb = INK
    callout.paragraph_format.space_after = Pt(0)
    callout.paragraph_format.line_spacing = 1.15

    header = section.header.paragraphs[0]
    header.text = "TPO Ciencia de Datos - Analisis de comportamiento"
    header.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    header.runs[0].font.size = Pt(9)
    header.runs[0].font.color.rgb = MUTED

    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    footer.add_run("Informe generado a partir del pipeline reproducible del proyecto")
    footer.runs[0].font.size = Pt(9)
    footer.runs[0].font.color.rgb = MUTED


def add_cover(doc: Document) -> None:
    doc.add_paragraph()
    doc.add_paragraph()
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("Analisis de comportamiento alrededor de parciales")
    run.bold = True
    run.font.size = Pt(24)
    run.font.color.rgb = BLUE

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = subtitle.add_run("Trabajo Practico - Ciencia de Datos")
    run.font.size = Pt(15)
    run.font.color.rgb = DARK_BLUE

    meta = doc.add_paragraph()
    meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = meta.add_run("Netflix, Spotify y datos de actividad fisica")
    run.font.size = Pt(11)
    run.font.color.rgb = MUTED

    doc.add_paragraph()
    add_callout(
        doc,
        "Resumen.",
        "El informe documenta que datos se integraron, como se procesaron, que metricas se calcularon y que resultados se observaron alrededor de fechas de parciales.",
    )

    doc.add_paragraph()
    cover_table = add_table(
        doc,
        ["Elemento", "Detalle"],
        [
            ["Archivo principal", "analisis_comportamiento_parciales.py"],
            ["Calendario", "cronograma_cuatrimestre_2026.csv"],
            ["Ventana de analisis", "14 dias antes y 14 dias despues de cada parcial"],
            ["Salida principal", "outputs/before_after_summary.csv"],
            ["Metodos de clustering", "K-Means y jerarquico aglomerativo"],
        ],
        [2600, 6760],
    )
    cover_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    doc.add_page_break()


def build_report() -> None:
    doc = Document()
    configure_document(doc)
    add_cover(doc)

    ranking = read_csv(BASE_DIR / "outputs" / "ranking_support_score.csv")
    top_ranking = ranking[:6]
    kmeans = {
        "Netflix": read_csv(output_csv("kmeans_evaluation_views.csv")),
        "Spotify": read_csv(output_csv("kmeans_evaluation_minutes.csv")),
        "Steps": read_csv(output_csv("kmeans_evaluation_steps.csv")),
    }
    agg = {
        "Netflix": read_csv(output_csv("agglomerative_evaluation_views.csv")),
        "Spotify": read_csv(output_csv("agglomerative_evaluation_minutes.csv")),
        "Steps": read_csv(output_csv("agglomerative_evaluation_steps.csv")),
    }

    doc.add_heading("1. Introduccion", level=1)
    doc.add_paragraph(
        "El proyecto analiza si los estudiantes modifican su comportamiento cerca de los parciales. Para responder esa pregunta se combinaron datos de consumo digital y actividad fisica, se agregaron por dia y se compararon ventanas antes y despues de cada examen."
    )
    doc.add_paragraph(
        "La propuesta es exploratoria: no se busca predecir notas ni establecer causalidad, sino encontrar senales consistentes que puedan sostener o matizar la hipotesis de cambio de rutina."
    )

    doc.add_heading("2. Fuentes de datos utilizadas", level=1)
    add_table(
        doc,
        ["Fuente", "Origen", "Variable analizada"],
        [
            ["Netflix", "Historiales de visualizacion CSV", "Cantidad de visualizaciones por dia"],
            ["Spotify", "StreamingHistory en JSON", "Minutos escuchados por dia"],
            ["Samsung Health", "Exportaciones de salud", "Pasos y registros de sueno"],
            ["Apple Health", "Export XML", "Pasos y registros de sueno"],
            ["Google Fit", "Takeout de metricas diarias", "Pasos diarios"],
            ["Calendario academico", "cronograma_cuatrimestre_2026.csv", "Fechas de parciales"],
        ],
        [1600, 3700, 4060],
    )

    doc.add_heading("3. Metodologia", level=1)
    doc.add_heading("3.1 Ventana temporal", level=2)
    doc.add_paragraph(
        "Cada parcial se tomo como evento central. Alrededor de esa fecha se construyo una ventana de -14 a +14 dias. Esto permite alinear series de distintas personas y materias usando una misma referencia temporal."
    )
    add_table(
        doc,
        ["Tramo", "Dias relativos", "Uso en el analisis"],
        [
            ["Baseline", "-14 a -8", "Referencia previa al periodo mas cercano al parcial"],
            ["Pre parcial", "-7 a -1", "Comportamiento antes del examen"],
            ["Dia del parcial", "0", "Valor observado el dia del evento"],
            ["Post parcial", "+1 a +7", "Comportamiento despues del examen"],
        ],
        [2100, 1900, 5360],
    )

    doc.add_heading("3.2 Transformacion y agregacion", level=2)
    add_number(doc, "Se leyeron las fuentes originales y se normalizaron a registros diarios por estudiante.")
    add_number(doc, "Se filtraron solo las fechas cercanas a parciales para reducir ruido y volumen.")
    add_number(doc, "Se calcularon promedios pre y post parcial para cada estudiante y fuente.")
    add_number(doc, "Se generaron rankings y visualizaciones para auditar los cambios mas relevantes.")

    doc.add_heading("3.3 Variables calculadas", level=2)
    add_table(
        doc,
        ["Variable", "Significado"],
        [
            ["pre_mean", "Promedio durante los 7 dias previos al parcial"],
            ["post_mean", "Promedio durante los 7 dias posteriores"],
            ["delta_abs", "Diferencia absoluta entre post y pre"],
            ["delta_pct", "Cambio porcentual respecto del tramo pre"],
            ["effect_size", "Cambio estandarizado segun la dispersion observada"],
            ["support_score", "Score heuristico que pondera magnitud y cantidad de datos"],
        ],
        [2200, 7160],
    )

    doc.add_heading("4. Resultados principales", level=1)
    add_callout(
        doc,
        "Hallazgo principal.",
        "La senal mas clara aparece en actividad fisica. Los cambios de pasos tienen mayor soporte que los cambios observados en Spotify o Netflix.",
    )
    add_table(
        doc,
        ["Estudiante", "Fuente", "Pre n", "Post n", "Delta %", "Support"],
        [
            [
                row.get("student", ""),
                row.get("source", ""),
                row.get("pre_n", ""),
                row.get("post_n", ""),
                fmt_num(row.get("delta_pct"), 2),
                fmt_num(row.get("support_score"), 2),
            ]
            for row in top_ranking
        ],
        [1500, 1300, 1000, 1000, 1300, 3260],
    )
    doc.add_paragraph(
        "La lectura de estos resultados indica que los cambios mas respaldados por volumen de datos aparecen en pasos. Algunos estudiantes aumentan la actividad despues del parcial, mientras que otros la reducen. Por lo tanto, la hipotesis se sostiene de manera parcial: hay cambios, pero no todos tienen la misma direccion."
    )

    doc.add_heading("5. Clustering", level=1)
    doc.add_paragraph(
        "Para explorar perfiles de comportamiento se represento a cada estudiante como un vector de valores diarios alrededor del parcial. Luego se estandarizaron las series y se aplicaron metodos de clustering no supervisado."
    )

    doc.add_heading("5.1 K-Means", level=2)
    doc.add_paragraph(
        "Se evaluaron distintos valores de k. Para cada opcion se calculo la inercia, usada como apoyo para el metodo del codo, y la silueta, usada para seleccionar el k con mejor separacion valida."
    )
    add_table(
        doc,
        ["Fuente", "k elegido", "Silueta"],
        [
            [source, selected_value(rows, "k"), fmt_num(selected_value(rows, "silhouette"), 3)]
            for source, rows in kmeans.items()
        ],
        [2500, 1800, 5060],
    )

    doc.add_heading("5.2 HDBSCAN", level=2)
    doc.add_paragraph(
        "Tambien se probo HDBSCAN como metodo basado en densidad. En esta muestra todos los puntos quedaban etiquetados como ruido (-1), por lo que el metodo no aportaba una segmentacion interpretable. Por ese motivo se saco del flujo principal."
    )

    doc.add_heading("5.3 Clustering jerarquico aglomerativo", level=2)
    doc.add_paragraph(
        "Luego se probo clustering jerarquico aglomerativo con enlace promedio. El resultado fue muy parecido al de K-Means, lo que refuerza que la estructura encontrada no depende exclusivamente de un unico algoritmo."
    )
    add_table(
        doc,
        ["Fuente", "k elegido", "Silueta", "Linkage"],
        [
            [
                source,
                selected_value(rows, "k"),
                fmt_num(selected_value(rows, "silhouette"), 3),
                selected_value(rows, "linkage") or "average",
            ]
            for source, rows in agg.items()
        ],
        [2200, 1600, 1600, 3960],
    )

    doc.add_heading("6. Visualizaciones y archivos generados", level=1)
    add_table(
        doc,
        ["Archivo", "Uso"],
        [
            ["before_after_summary.csv", "Resumen pre/post por estudiante y fuente"],
            ["ranking_support_score.csv", "Ranking recomendado para priorizar evidencia"],
            ["*_relative_day.svg", "Tendencias por dia relativo al parcial"],
            ["*_heatmap.svg", "Heterogeneidad por estudiante y dia relativo"],
            ["clusters_*_comparison.csv", "Comparacion entre K-Means y jerarquico"],
            ["interactive_dashboard.html", "Dashboard con filtros y vista exploratoria"],
        ],
        [3100, 6260],
    )

    doc.add_heading("7. Limitaciones", level=1)
    add_bullet(doc, "La muestra es chica, por lo que los clusters deben interpretarse como exploratorios.")
    add_bullet(doc, "Las plataformas de salud pueden medir actividad de manera distinta segun el dispositivo.")
    add_bullet(doc, "Algunas fuentes tienen mas cobertura temporal que otras.")
    add_bullet(doc, "Los cambios observados son asociaciones temporales, no evidencia causal.")
    add_bullet(doc, "La silueta puede favorecer clusters muy chicos cuando hay pocos estudiantes.")

    doc.add_heading("8. Conclusion", level=1)
    doc.add_paragraph(
        "La evidencia sugiere que existen cambios de comportamiento alrededor de los parciales, especialmente en actividad fisica medida por pasos. La direccion del cambio no es igual para todos: algunos estudiantes aumentan su actividad despues del examen y otros la reducen."
    )
    doc.add_paragraph(
        "Spotify muestra variaciones mas puntuales y Netflix funciona mejor como evidencia complementaria. Los metodos de clustering ayudan a explorar perfiles, pero no deben presentarse como clasificaciones definitivas por el tamano reducido de la muestra."
    )
    add_callout(
        doc,
        "Lectura para defensa.",
        "El argumento mas solido es que la rutina fisica cambia cerca de parciales; el consumo digital tambien puede variar, pero con menor consistencia entre estudiantes.",
    )

    doc.save(OUTPUT)


def selected_value(rows: list[dict[str, str]], key: str) -> str:
    for row in rows:
        if str(row.get("selected", "")).lower() == "true":
            return row.get(key, "")
    return "-"


if __name__ == "__main__":
    build_report()
